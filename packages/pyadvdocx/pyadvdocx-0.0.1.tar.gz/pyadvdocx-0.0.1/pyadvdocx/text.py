from docx import Document

def replace(docx:Document,**kwargs):
    """
        A function that replaces the given key word args from the given document.

        :param docx Document object from python-docx package
        :param kwargs Dictionary of keyword arguments that needs to be replaced
    """

    # Get a dict of the kwargs passed to replace
    __replace_with={}
    __replace_with.update(kwargs)

    __replace_in_header_and_footer(docx, __replace_with)
    __replace_in_tables(docx.tables, __replace_with)

    # Do replace in body
    __replace_in_paragraphs(docx.paragraphs,__replace_with)
        

def __replace_in_header_and_footer(docx:Document,__replace_with:dict):
    for section in docx.sections:
        header=section.header
        footer=section.footer
        __replace_in_paragraphs(header.paragraphs,__replace_with)
        __replace_in_paragraphs(footer.paragraphs,__replace_with)


def __replace_in_tables(tables,__replace_with:dict):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                if len(cell.tables)>0:
                    __replace_in_tables(cell.tables,__replace_with)
                if len(cell.paragraphs)>0:
                    __replace_in_paragraphs(cell.paragraphs,__replace_with)

                if len(cell.text)>0:
                    new_text=""
                    for word in cell.text.split(" "):
                        if word in __replace_with:
                            word=__replace_with[word]
                        new_text+=word
                    cell.text=new_text

def __replace_in_paragraphs(paragraphs,__replace_with):
    for p in paragraphs:
        new_text=""
        for word in p.text.split(" "):
            if word in __replace_with:
                word=__replace_with[word]
            new_text+=word
        p.text=new_text  

