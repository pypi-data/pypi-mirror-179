import re
from io import BytesIO
import json
from docx import Document
import importlib.resources


class Generator:
    def __init__(self):
        self.configuration = self.get_default_config()
        self.document = None
        self.previous_indented = None
        self.verse_count = None

    @staticmethod
    def get_default_config():
        return {"style_mappings": {"number": "number", "verse": "verse", "chorus": "chorus", "title": "title-itself",
                                   "footer": "prayer"},
                "template": "template.docx",
                "footer": "prayer.json"}

    def generate(self, songbook, target_stream=None):
        if target_stream is None:
            target_stream = BytesIO()
        self.document = Document(importlib.resources.files("resources").joinpath(self.configuration['template']).open('rb'))
        for song in songbook.songs:
            self.add_song_to_document(song)
        self.add_footer()
        self.document.save(target_stream)
        return target_stream

    def add_footer(self):
        style_name = self.configuration['style_mappings']['footer']
        try:
            file = importlib.resources.files("resources").joinpath(self.configuration['footer']).open()
            content = json.load(file)['content']
            self.document.add_paragraph(content, style=style_name)
        except FileNotFoundError:
            pass

    def add_song_to_document(self, song):
        self.verse_count = 0
        self.previous_indented = True
        self.document.add_paragraph('', style=self.configuration['style_mappings']['number'])
        self.handle_title(song)
        songparts = song.content
        for songpart in songparts[1:]:
            self.add_songpart_to_document(songpart)

    def handle_title(self, song):
        pattern = song.title
        first_part = song.content[0]
        embedded_title = self.find_title_in_songpart(song.title, first_part)
        if embedded_title is None:
            self.add_title(song.title)
            self.add_songpart_to_document(first_part)
        else:
            self.add_songpart_including_title(first_part, embedded_title)

    def find_title_in_songpart(self, title, songpart):
        pattern = title
        match = re.search(pattern, songpart.lyrics())
        if match is None:
            return None
        return match.group()

    def add_title(self, title):
        style_name = self.configuration['style_mappings']['title']
        self.document.add_paragraph(title, style=style_name)

    def add_songpart_including_title(self, songpart, title):
        self.previous_part_type = songpart.type
        style_name = self.configuration['style_mappings'][songpart.type]
        if not self.previous_indented:
            style_name += '-indent'
        self.previous_indented = not self.previous_indented
        lines = songpart.lyrics().split('\n')
        first_line_spaced = " " + lines[0].strip() + " "
        first_line_no_title = first_line_spaced.split(title)
        paragraph = self.document.add_paragraph('', style=style_name)
        prefix = ""
        if songpart.marked:
            if songpart.type == "verse":
                self.verse_count += 1
                prefix = str(self.verse_count) + ". "
            elif songpart.type == "chorus":
                prefix = "Ref.: "
        pre_title = prefix + first_line_no_title[0]
        if pre_title != ' ':
            paragraph.add_run(pre_title.lstrip())
        title_run = paragraph.add_run(title)
        title_run.style = 'title-' + style_name
        post_title = first_line_no_title[1]
        if post_title != ' ':
            paragraph.add_run(post_title.rstrip())
        for line in lines[1:]:
            paragraph.add_run("" + '\n' + line)

    def add_songpart_to_document(self, songpart):
        self.previous_part_type = songpart.type
        style_name = self.configuration['style_mappings'][songpart.type]
        if not self.previous_indented:
            style_name += '-indent'
        self.previous_indented = not self.previous_indented
        prefix = ""
        if songpart.marked:
            if songpart.type == "verse":
                self.verse_count += 1
                prefix = str(self.verse_count) + ". "
            elif songpart.type == "chorus":
                prefix = "Ref.: "
        self.document.add_paragraph(prefix + songpart.lyrics(), style=style_name)
